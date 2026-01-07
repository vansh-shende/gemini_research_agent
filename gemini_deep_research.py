import streamlit as st
from google import genai
import traceback
import io
try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    Document = None
    DOCX_AVAILABLE = False

# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="Gemini Research Assistant",
    page_icon="🔍",
    layout="wide"
)

st.title("🔍 Gemini Research Assistant")
st.markdown("Research powered by **Google Gemini API**")

# ---------------- SIDEBAR ----------------
st.sidebar.header("🔑 API Key")
api_key = st.sidebar.text_input("Enter Gemini API Key", type="password")

# ---------------- MODEL LISTING & SELECTION ----------------
def list_models(api_key):
    """Return a list of available model names using the provided API key.

    Tries several SDK call shapes, surfaces errors in the sidebar, and parses
    different response shapes.
    """
    try:
        client = genai.Client(api_key=api_key)
    except Exception as e:
        st.sidebar.error(f"Client init error: {e}")
        return []

    resp = None
    errors = []

    # Try common call shapes; prefer client.models.list() which returns a pager
    try_calls = [
        ("client.models.list()", lambda c: c.models.list() if hasattr(c, "models") and hasattr(c.models, "list") else None),
        ("client.list_models()", lambda c: c.list_models() if hasattr(c, "list_models") else None),
        ("client.list()", lambda c: c.list() if hasattr(c, "list") else None),
    ]

    for name, fn in try_calls:
        try:
            candidate = fn(client)
            if candidate is not None:
                resp = candidate
                break
        except Exception as e:
            errors.append(f"{name} -> {e}")

    if resp is None:
        st.sidebar.error("Could not list models. See attempts below:")
        for e in errors:
            st.sidebar.text(e)
        return []

    # Normalize response to a list of model dicts/objects or iterate pager
    models = []

    # If resp is a dict-like container
    if isinstance(resp, dict):
        candidate_list = resp.get("models") or resp.get("model") or resp.get("data") or []
        for m in candidate_list or []:
            name = m.get("name") or m.get("id") if isinstance(m, dict) else getattr(m, "name", None) or getattr(m, "id", None)
            if name:
                models.append(name)

    # If resp is a list or similar
    elif isinstance(resp, list):
        for m in resp:
            name = m.get("name") if isinstance(m, dict) else getattr(m, "name", None) or getattr(m, "id", None)
            if name:
                models.append(name)

    else:
        # Attempt to iterate (e.g., Pager object) and extract model names
        try:
            for m in resp:
                name = None
                if isinstance(m, dict):
                    name = m.get("name") or m.get("id")
                else:
                    # model objects often have 'name' or 'id'
                    name = getattr(m, "name", None) or getattr(m, "id", None)
                if name:
                    models.append(name)
        except Exception:
            # Fall back to inspecting attributes
            candidate_list = getattr(resp, "models", None) or getattr(resp, "data", None) or getattr(resp, "model", None) or []
            for m in candidate_list or []:
                name = m.get("name") if isinstance(m, dict) else getattr(m, "name", None) or getattr(m, "id", None)
                if name:
                    models.append(name)

    if not models:
        st.sidebar.info("No models returned by the API. Check that your key has Gemini Pro access and call ListModels output for supported methods.")

    return sorted(set(models))


# Attempt to populate model selector if API key is provided
available_models = []
if api_key:
    with st.spinner("Loading available models..."):
        available_models = list_models(api_key)

# Prefer a sensible default if present
if available_models:
    try:
        default_index = available_models.index("models/gemini-pro-latest")
    except ValueError:
        default_index = 0
    model_choice = st.sidebar.selectbox("Choose model", options=available_models, index=default_index)
else:
    model_choice = st.sidebar.selectbox("Choose model", options=["(enter API key to load models)"])

st.sidebar.markdown("**Tip:** If you get 404 NOT_FOUND, pick a model from the list above (e.g., `models/gemini-pro-latest`) — `gemini-1.5-pro-latest` may not exist for this API/version.")

# ---------------- DEBUG / DIAGNOSTICS ----------------
def list_models_debug(api_key):
    """Run multiple ListModels call shapes and return raw summaries and any errors."""
    try:
        client = genai.Client(api_key=api_key)
    except Exception as e:
        return {"error": f"Client init error: {e}", "traceback": traceback.format_exc()}

    attempts = []

    try_calls = [
        ("client.list_models()", lambda c: c.list_models() if hasattr(c, "list_models") else None),
        ("client.list()", lambda c: c.list() if hasattr(c, "list") else None),
        ("client.models.list()", lambda c: c.models.list() if hasattr(c, "models") and hasattr(c.models, "list") else None),
        ("client.models.list_models()", lambda c: c.models.list_models() if hasattr(c, "models") and hasattr(c.models, "list_models") else None),
    ]

    for name, fn in try_calls:
        try:
            candidate = fn(client)
            # Summarize candidate safely
            try:
                if isinstance(candidate, (list, tuple)):
                    summary = {"type": str(type(candidate)), "len": len(candidate), "sample": repr(candidate[:5])}
                elif isinstance(candidate, dict):
                    summary = {"type": "dict", "keys": list(candidate.keys())}
                else:
                    summary = {"type": str(type(candidate)), "repr": repr(candidate)[:1000]}
            except Exception as e2:
                summary = {"type": str(type(candidate)), "repr_error": str(e2)}

            attempts.append({"name": name, "ok": True, "summary": summary})
        except Exception as e:
            attempts.append({"name": name, "ok": False, "error": str(e), "traceback": traceback.format_exc()})

    models = list_models(api_key)
    return {"attempts": attempts, "models": models}


with st.expander("🔎 Debug model listing"):
    if not api_key:
        st.info("Enter API key to run model listing debug")
    else:
        debug = list_models_debug(api_key)
        if debug.get("error"):
            st.error(debug["error"])
            st.text(debug["traceback"])
        else:
            for a in debug.get("attempts", []):
                if a.get("ok"):
                    st.success(a.get("name"))
                    st.json(a.get("summary"))
                else:
                    st.error(a.get("name"))
                    st.text(a.get("error"))
                    st.text(a.get("traceback"))

            st.write("Parsed models:", debug.get("models"))

# ---------------- USER INPUT ----------------
topic = st.text_input(
    "Enter your research topic",
    placeholder="e.g. Future of AI in Healthcare"
)

# ---------------- GEMINI FUNCTION ----------------
def run_gemini(api_key, topic, model_name):
    client = genai.Client(api_key=api_key)

    prompt = f"""
You are a professional academic research assistant.

Write a detailed research report on:
"{topic}"

Include:
1. Introduction
2. Background
3. Key Concepts
4. Real-world Examples
5. Advantages & Limitations
6. Future Scope
7. Conclusion

Use simple English.
"""

    # Primary attempt: use the SDK's generate_content if available for the model
    try:
        response = client.models.generate_content(
            model=model_name,
            contents=prompt
        )
        # Many SDK shapes expose `.text` or `.content` — be defensive
        return getattr(response, "text", None) or getattr(response, "content", None) or str(response)

    except Exception as e:
        # If model isn't available for generate_content, give a clear error message.
        msg = str(e)
        if "not found" in msg.lower() or "not supported for generatecontent" in msg.lower():
            raise RuntimeError(
                "Selected model does not support `generate_content` for this API/version. "
                "Call ListModels and choose a model that supports text generation or use a different SDK method."
            )
        if "resource_exhausted" in msg.lower() or "quota" in msg.lower() or "429" in msg:
            raise RuntimeError(
                "Quota exhausted or insufficient billing for this model. "
                "Enable billing or request quota for generative requests at https://ai.google.dev/gemini-api/docs/rate-limits and monitor usage at https://ai.dev/usage?tab=rate-limit."
            )
        # Re-raise other unexpected exceptions with traceback to help debugging
        tb = traceback.format_exc()
        raise RuntimeError(f"Unexpected error calling model: {e}\n\n{tb}")

# ---------------- WORD EXPORT ----------------
def make_word_document(topic, content):
    """Return a BytesIO containing a .docx file for the generated content."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
    doc = Document()
    doc.add_heading(topic, level=1)
    # Split into paragraphs by double-newline to keep structure
    for block in content.split("\n\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ---------------- BUTTON ----------------
if st.button("🚀 Start Research"):
    if not api_key:
        st.warning("Please enter API key")
    elif not topic:
        st.warning("Please enter topic")
    elif not model_choice or model_choice.startswith("("):
        st.warning("Please choose a valid model from the sidebar")
    else:
        with st.spinner("Generating research..."):
            try:
                output = run_gemini(api_key, topic, model_choice)

                st.subheader("📘 Research Output")
                st.write(output)

                st.download_button(
                    "⬇️ Download Report",
                    output,
                    file_name=f"{topic.replace(' ', '_')}.txt"
                )

                # DOCX download button
                try:
                    docx_bio = make_word_document(topic, output)
                    st.download_button(
                        "⬇️ Download as Word (.docx)",
                        data=docx_bio.getvalue(),
                        file_name=f"{topic.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.info("Enable Word export: run `pip install python-docx` in your venv")

            except Exception as e:
                st.error(f"❌ Error: {e}")

# ---------------- FOOTER ----------------
st.markdown("---")
st.caption("Powered by Google Gemini 🚀")
